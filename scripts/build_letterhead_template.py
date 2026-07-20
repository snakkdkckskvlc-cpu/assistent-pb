#!/usr/bin/env python3
"""Собирает resources/templates/letterhead.docx из канонического бланка компании.

Источник правды — файл «бланк ПожСервис новый.docx» (обычно лежит в личной
папке компании, вне репозитория — коммерческий документ). Этот скрипт берёт
его КАК ЕСТЬ и точечно вставляет 5 плейсхолдеров в уже существующие абзацы/
ячейки бланка — не добавляет и не удаляет ни одного абзаца, чтобы результат
оставался структурно идентичен оригиналу (тот же шрифт, отступы, пустые
строки под текст письма).

Куда вставляются плейсхолдеры (см. infrastructure/generators/letter_docx.py):
  - "Уважаемый (-ая) ______!" → {{greeting}}
  - первая строка тела письма (с табуляцией-отступом) → {{body}}
    (эта же строка сохраняет табуляцию — {{body}} подставляется ПОСЛЕ неё;
    многоабзацный текст письма при подстановке добавляет новые абзацы сразу
    после этого, не трогая остальные пустые строки бланка)
  - строка подписи "Директор ... Фамилия И.О." → {{sender_position}} / {{sender_name}}
  - в таблице с реквизитами: дата, тема письма, получатель

Запуск:
    python scripts/build_letterhead_template.py --source "/путь/к/бланк ПожСервис новый.docx"

По умолчанию источник ищется в стандартном месте, куда его обычно кладут
(см. --source по умолчанию ниже) и результат пишется поверх
apps/backend/src/fire_safety_backend/resources/templates/letterhead.docx.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT = (
    ROOT
    / "apps"
    / "backend"
    / "src"
    / "fire_safety_backend"
    / "resources"
    / "templates"
    / "letterhead.docx"
)

_GREETING_MARKER = "Уважаемый"
_SIGNOFF_MARKER = "Директор"
_DATE_MARKER = "№"
_SUBJECT_MARKER = "тема письма"
_RECIPIENT_MARKER = "Руководителю"


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Полностью заменяет видимый текст абзаца одним run'ом — форматирование
    (шрифт/размер) берёт у первого run'а исходного абзаца, остальные run'ы
    (если текст был раздёлен на несколько кусков форматированием) очищает."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_signoff_line(paragraph, position_tag: str, name_tag: str) -> None:
    """Строка подписи — два значимых run'а (должность, ФИО) с табуляцией
    между ними для выравнивания; табы должны остаться нетронутыми."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if len(runs) < 2:
        raise ValueError(f"Не нашёл в строке подписи два непустых run'а: {paragraph.text!r}")
    runs[0].text = position_tag
    runs[-1].text = name_tag


def build(source: Path, output: Path) -> None:
    from docx import Document

    doc = Document(str(source))

    paragraphs = doc.paragraphs
    greeting_idx = next((i for i, p in enumerate(paragraphs) if _GREETING_MARKER in p.text), None)
    if greeting_idx is None:
        raise ValueError(f"Не нашёл абзац с обращением («{_GREETING_MARKER}») в {source}")
    _replace_paragraph_text(paragraphs[greeting_idx], "{{greeting}}")

    # Первый абзац тела письма — сразу после приветствия, обычно это
    # табуляция (отступ первой строки), которую сохраняем перед плейсхолдером.
    body_para = paragraphs[greeting_idx + 1]
    _replace_paragraph_text(body_para, body_para.text + "{{body}}")

    signoff_idx = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip().startswith(_SIGNOFF_MARKER)), None
    )
    if signoff_idx is None:
        raise ValueError(f"Не нашёл строку подписи («{_SIGNOFF_MARKER}...») в {source}")
    _replace_signoff_line(paragraphs[signoff_idx], "{{sender_position}}", "{{sender_name}}")

    # Таблица реквизитов: дата/номер письма, тема, получатель.
    if not doc.tables:
        raise ValueError(f"В {source} нет таблицы с реквизитами")
    table = doc.tables[0]
    date_done = subject_done = recipient_done = False
    for row in table.rows:
        for cell in row.cells:
            cell_paragraphs = cell.paragraphs
            for pi, p in enumerate(cell_paragraphs):
                text = p.text
                if not date_done and text.strip().startswith("_") and _DATE_MARKER in text:
                    _replace_paragraph_text(p, "{{date}}" + text[text.index(_DATE_MARKER) - 1 :])
                    date_done = True
                elif not subject_done and _SUBJECT_MARKER in text:
                    _replace_paragraph_text(p, "{{subject}}")
                    subject_done = True
                elif not recipient_done and text.strip().startswith(_RECIPIENT_MARKER):
                    _replace_paragraph_text(p, "{{recipient}}")
                    recipient_done = True
                    # Остальные абзацы этой ячейки — пример ФИО получателя
                    # («Иванову А.А.» и т.п.), часть примера в бланке.
                    # {{recipient}} сам многострочный (подстановка добавляет
                    # нужные абзацы), а статичный текст-пример надо стереть —
                    # иначе он попадёт в каждое сгенерированное письмо.
                    for later in cell_paragraphs[pi + 1 :]:
                        if later.text.strip():
                            _replace_paragraph_text(later, "")
    missing = [
        name
        for name, done in [
            ("date", date_done),
            ("subject", subject_done),
            ("recipient", recipient_done),
        ]
        if not done
    ]
    if missing:
        raise ValueError(f"Не нашёл в таблице реквизитов ячейки для: {', '.join(missing)}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        "--source", required=True, type=Path, help="Путь к каноническому бланку (.docx)"
    )
    parser.add_argument("--out", type=Path, default=DEFAULT_OUTPUT, help="Куда сохранить результат")
    args = parser.parse_args()

    if not args.source.exists():
        print(f"Не найден исходный файл: {args.source}", file=sys.stderr)
        return 1

    build(args.source, args.out)
    print(f"Готово: {args.out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
