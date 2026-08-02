"""Юнит-тесты infrastructure/generators/letter_docx.py.

resources/templates/letterhead.docx поставляется вместе с кодом (репозиторий
приватный), поэтому на обычной установке доступны оба уровня тестов:
  - тесты против РЕАЛЬНОГО шаблона — основной путь, пропускаются только если
    файл повреждён или удалён;
  - тесты против запасного пути (шаблона нет вовсе) — он остаётся для
    повреждённой установки и проверяется подменой пути, а не отсутствием
    файла в репозитории.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fire_safety_backend import config
from fire_safety_backend.infrastructure.generators.letter_docx import build_letter_docx

_EXPECTED_PLACEHOLDERS = {
    "date",
    "recipient",
    "subject",
    "greeting",
    "body",
    "sender_position",
    "sender_name",
}

_HAS_REAL_TEMPLATE = config.LETTERHEAD_TEMPLATE.exists()
_SKIP_REASON = (
    "letterhead.docx отсутствует. Обычно он поставляется вместе с кодом "
    "(репозиторий приватный), так что пропуск здесь означает повреждённую "
    "установку. Восстановить: git checkout -- "
    "apps/backend/src/fire_safety_backend/resources/templates/letterhead.docx"
)

_SAMPLE_LETTER = {
    "получатель": "Директору\nООО «Ромашка»\n\nИванову И.И.",
    "тема": "О проведении планового ТО",
    "обращение": "Уважаемый Иван Иванович!",
    "тело": "Первый абзац письма.\n\nВторой абзац письма.",
    "должность_отправителя_placeholder": "Директор",
    "фио_отправителя_placeholder": "О.Н. Сляднев",
}


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# --- Против реального (локального, коммерческого) шаблона — пропускаются в CI ---


@pytest.mark.skipif(not _HAS_REAL_TEMPLATE, reason=_SKIP_REASON)
def test_committed_template_has_exactly_expected_placeholders() -> None:
    doc = Document(str(config.LETTERHEAD_TEMPLATE))
    text = _all_text(doc)
    for tag in _EXPECTED_PLACEHOLDERS:
        assert f"{{{{{tag}}}}}" in text, f"Плейсхолдер {{{{{tag}}}}} не найден в шаблоне"
    # Старое поле "формула_вежливости"/{{signoff}} убрано — в реальном бланке
    # для него нет отдельной строки (см. scripts/build_letterhead_template.py).
    assert "{{signoff}}" not in text


@pytest.mark.skipif(not _HAS_REAL_TEMPLATE, reason=_SKIP_REASON)
def test_template_has_no_leftover_example_text() -> None:
    # Регрессия: билдер шаблона однажды оставил "Иванову А.А." как статичный
    # текст рядом с {{recipient}} — пример из исходного бланка, который
    # попадал бы в КАЖДОЕ сгенерированное письмо.
    doc = Document(str(config.LETTERHEAD_TEMPLATE))
    text = _all_text(doc)
    assert "Иванову" not in text
    assert "Уважаемый (-ая)" not in text


@pytest.mark.skipif(not _HAS_REAL_TEMPLATE, reason=_SKIP_REASON)
def test_build_letter_docx_against_real_template(tmp_path: Path) -> None:
    out = tmp_path / "letter.docx"
    build_letter_docx(_SAMPLE_LETTER, out)
    doc = Document(str(out))
    text = _all_text(doc)

    assert "{{" not in text, f"Незамещённый плейсхолдер остался в выводе: {text!r}"
    assert "Уважаемый Иван Иванович!" in text
    assert "Первый абзац письма." in text
    assert "Второй абзац письма." in text
    assert "О проведении планового ТО" in text
    assert "Директору" in text
    assert "ООО «Ромашка»" in text
    assert "Иванову И.И." in text
    assert "О.Н. Сляднев" in text


# --- Против fallback-пути (шаблона нет) — гарантированно прогоняются всюду ---


def test_build_letter_docx_fallback_when_template_missing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Запасной путь build_letter_docx: бланка нет — собираем простой DOCX,
    а не падаем.

    Бланк теперь поставляется с кодом, поэтому в обычной установке этот путь
    не выполняется. Он остаётся для повреждённой установки и для случая, когда
    файл ещё не восстановлен, — и проверяется здесь подменой пути, а не
    отсутствием файла в репозитории."""
    monkeypatch.setattr(config, "LETTERHEAD_TEMPLATE", tmp_path / "does-not-exist.docx")
    out = tmp_path / "out.docx"

    build_letter_docx(_SAMPLE_LETTER, out)

    assert out.exists()
    doc = Document(str(out))
    text = _all_text(doc)
    assert "{{" not in text
    assert "Уважаемый Иван Иванович!" in text
    assert "Первый абзац письма." in text
    assert "Директору" in text


def test_build_letter_docx_fallback_handles_empty_fields(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(config, "LETTERHEAD_TEMPLATE", tmp_path / "does-not-exist.docx")
    out = tmp_path / "empty.docx"

    build_letter_docx({}, out)

    doc = Document(str(out))
    assert "{{" not in _all_text(doc)
