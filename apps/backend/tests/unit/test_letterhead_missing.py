"""Пропавший бланк обязан быть заметен ДО отправки письма.

Реальный случай: шаблон letterhead.docx исчез (он в .gitignore и в git
никогда не попадал), генератор молча свалился на запасной путь и стал
выдавать письма БЕЗ реквизитов, ИНН и банковских данных. Файл при этом
создавался, скачивался и открывался — разницу заметили только сравнив два
письма вручную. Для письма контрагенту это отправка неофициального
документа под видом официального.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fire_safety_backend import config
from fire_safety_backend.infrastructure.generators.letter_docx import build_letter_docx

_FIELDS = {
    "получатель": "Руководителю организации",
    "тема": "О выполненных работах",
    "обращение": "Уважаемый Иван Иванович!",
    "тело": "Работы выполнены в полном объёме.",
    "должность_отправителя_placeholder": "Директор",
    "фио_отправителя_placeholder": "О.Н. Сляднев",
}


def test_missing_template_is_logged_as_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    """Не warning: письмо без реквизитов нельзя отправлять, это отказ."""
    monkeypatch.setattr(config, "LETTERHEAD_TEMPLATE", tmp_path / "нет-такого.docx")
    with caplog.at_level("ERROR"):
        build_letter_docx(_FIELDS, tmp_path / "out.docx")

    errors = [r for r in caplog.records if r.levelname == "ERROR"]
    assert errors, "пропажа бланка прошла без ошибки в логе"
    text = errors[0].getMessage()
    assert "бланк" in text.lower()
    # В сообщении должен быть способ починить, а не только констатация.
    assert "build_letterhead_template" in text


def test_fallback_still_produces_a_readable_letter(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Текст письма терять нельзя — пользователь его уже отредактировал."""
    monkeypatch.setattr(config, "LETTERHEAD_TEMPLATE", tmp_path / "нет-такого.docx")
    out = build_letter_docx(_FIELDS, tmp_path / "out.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Работы выполнены в полном объёме." in text
    assert "О.Н. Сляднев" in text


def test_fallback_letter_has_no_company_requisites(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Фиксируем ЦЕНУ отказа: именно этого в документе и не хватает.

    Тест существует, чтобы «запасной» путь не начали считать равноценным
    бланку — в нём нет ни ИНН, ни расчётного счёта.
    """
    monkeypatch.setattr(config, "LETTERHEAD_TEMPLATE", tmp_path / "нет-такого.docx")
    out = build_letter_docx(_FIELDS, tmp_path / "out.docx")
    doc = Document(str(out))
    everything = "\n".join(p.text for p in doc.paragraphs) + "".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    assert "4823026435" not in everything, "ИНН"
    assert "Сбербанк" not in everything, "банковские реквизиты"


@pytest.mark.skipif(
    not config.LETTERHEAD_TEMPLATE.exists(),
    reason="бланк не установлен на этой машине (он в .gitignore)",
)
def test_real_template_puts_requisites_into_letter(tmp_path: Path) -> None:
    """Когда бланк есть — реквизиты обязаны оказаться в документе.

    Ровно этим отличалось «правильное» письмо от «неправильного»: наличием
    таблицы с реквизитами.
    """
    out = build_letter_docx(_FIELDS, tmp_path / "out.docx")
    doc = Document(str(out))
    assert doc.tables, "в бланке реквизиты лежат в таблице — её нет"
    cells = "".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "4823026435" in cells, "ИНН не подставился"
    assert "Сбербанк" in cells, "банковские реквизиты не подставились"
