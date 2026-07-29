"""Исправленный документ = копия оригинала с правками внутри.

Смысл фичи именно в сохранении вёрстки: документ потом отправляют и
подписывают. Если бы мы собирали DOCX заново из исправленного текста, на
выходе была бы простыня без заголовков, таблиц и бланка — и её всё равно
пришлось бы переверстывать руками, то есть кнопка «скачать исправленный»
не имела бы смысла.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fire_safety_backend import config
from fire_safety_backend.infrastructure.generators.corrected_docx import (
    _replacements,
    build_corrected_docx,
)


@pytest.fixture(autouse=True)
def _isolated_output(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "out")
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


_ERRORS = [
    {"type": "орфография", "before": "ответственость", "after": "ответственность"},
    {"type": "грамматика", "before": "в течении", "after": "в течение"},
    {"type": "грамматика", "before": "согласно графика", "after": "согласно графику"},
]


def _source(tmp_path: Path) -> Path:
    path = tmp_path / "приказ.docx"
    doc = Document()
    doc.add_heading("ПРИКАЗ о противопожарном режиме", level=1)
    para = doc.add_paragraph("Срок — ")
    para.add_run("в течении").bold = True
    para.add_run(" трёх дней.")
    doc.add_paragraph("Несёт ответственость за эвакуацию.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Пункт"
    table.rows[0].cells[1].text = "Проверка проводится согласно графика."
    doc.save(path)
    return path


def test_corrections_applied_in_body(tmp_path: Path) -> None:
    out, is_copy = build_corrected_docx("неважно", _ERRORS, _source(tmp_path))
    assert is_copy is True
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "ответственность" in text
    assert "ответственость" not in text
    assert "в течение трёх дней" in text


def test_corrections_applied_inside_tables(tmp_path: Path) -> None:
    """В договорах и приказах часть текста живёт в таблицах."""
    out, _ = build_corrected_docx("неважно", _ERRORS, _source(tmp_path))
    cell = Document(str(out)).tables[0].rows[0].cells[1].text
    assert "согласно графику" in cell
    assert "согласно графика" not in cell


def test_heading_and_table_styles_survive(tmp_path: Path) -> None:
    """Ради этого фича и делается копией, а не новым файлом."""
    out, _ = build_corrected_docx("неважно", _ERRORS, _source(tmp_path))
    doc = Document(str(out))
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert len(doc.tables) == 1
    assert doc.tables[0].style.name == "Table Grid"


def test_bold_survives_when_correction_fits_one_run(tmp_path: Path) -> None:
    """Жирным выделяют важное — терять это на каждой правке нельзя.

    «в течении» лежит целиком внутри жирного run'а, поэтому правка обязана
    примениться БЕЗ слияния run'ов.
    """
    out, _ = build_corrected_docx("неважно", _ERRORS, _source(tmp_path))
    para = Document(str(out)).paragraphs[1]
    assert any(r.bold and "в течение" in (r.text or "") for r in para.runs)


def test_correction_split_across_runs_is_applied(tmp_path: Path) -> None:
    """Word дробит абзац по своей логике, разрывая слова между run'ами."""
    path = tmp_path / "split.docx"
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Несёт ответствен")
    para.add_run("ость за эвакуацию.")
    doc.save(path)

    out, is_copy = build_corrected_docx("неважно", _ERRORS, path)
    assert is_copy is True
    assert "ответственность" in Document(str(out)).paragraphs[0].text


def test_original_file_is_not_modified(tmp_path: Path) -> None:
    """Правим КОПИЮ. Затереть присланный документ — потеря данных."""
    src = _source(tmp_path)
    before = src.read_bytes()
    out, _ = build_corrected_docx("неважно", _ERRORS, src)
    assert out != src
    assert src.read_bytes() == before


def test_non_docx_source_falls_back_to_plain_document(tmp_path: Path) -> None:
    """PDF и сканы «поправить внутри» нельзя — но скачать что-то надо."""
    pdf = tmp_path / "скан.pdf"
    pdf.write_bytes(b"%PDF-1.4 not a real pdf")
    out, is_copy = build_corrected_docx("Исправленный текст документа.", _ERRORS, pdf)
    assert is_copy is False
    assert "Исправленный текст документа." in Document(str(out)).paragraphs[0].text


def test_pasted_text_without_source_still_gives_document(tmp_path: Path) -> None:
    out, is_copy = build_corrected_docx("Первая строка.\nВторая строка.", [], None)
    assert is_copy is False
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert "Первая строка." in texts
    assert "Вторая строка." in texts


def test_broken_docx_does_not_lose_the_result(tmp_path: Path) -> None:
    """Нестандартный DOCX не должен лишать пользователя исправленного текста."""
    fake = tmp_path / "битый.docx"
    fake.write_bytes(b"PK\x03\x04 not a real docx")
    out, is_copy = build_corrected_docx("Исправленный текст.", _ERRORS, fake)
    assert is_copy is False
    assert "Исправленный текст." in Document(str(out)).paragraphs[0].text


def test_document_without_errors_is_still_a_valid_copy(tmp_path: Path) -> None:
    out, is_copy = build_corrected_docx("неважно", [], _source(tmp_path))
    assert is_copy is True
    assert Document(str(out)).paragraphs[0].style.name == "Heading 1"


class TestReplacementPairs:
    def test_longer_before_goes_first(self) -> None:
        """Иначе короткая замена съест часть длинной и сломает её."""
        pairs = _replacements(
            [{"before": "в", "after": "во"}, {"before": "в течении", "after": "в течение"}]
        )
        assert pairs[0][0] == "в течении"

    def test_empty_and_noop_pairs_dropped(self) -> None:
        pairs = _replacements(
            [
                {"before": "", "after": "что-то"},
                {"before": "слово", "after": "слово"},
                {"before": "ошибка", "after": "исправление"},
            ]
        )
        assert pairs == [("ошибка", "исправление")]

    def test_duplicates_dropped(self) -> None:
        pairs = _replacements([{"before": "а", "after": "б"}, {"before": "а", "after": "б"}])
        assert pairs == [("а", "б")]

    def test_malformed_entries_do_not_crash(self) -> None:
        """Модель иногда отступает от схемы — список строк вместо объектов."""
        assert _replacements(["строка", None, 42, {"before": "а", "after": "б"}]) == [("а", "б")]
