"""Договор с правками в режиме рецензирования Word.

Смысл фичи в том, что контрагент открывает файл и жмёт «принять» или
«отклонить» по каждой правке. Поэтому проверяется не «в файле есть нужные
слова», а два состояния документа: что увидит принявший все правки и что
увидит отклонивший. Если хотя бы одно из них неверно, юрист отправит
контрагенту испорченный договор — а это дороже, чем не отправить ничего.

Отдельно проверяется структура разметки: Word молча выбрасывает правку, если
внутри w:del лежит w:t вместо w:delText, и файл открывается без правок вовсе.
Такую поломку по тексту не увидеть.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from fire_safety_backend import config
from fire_safety_backend.infrastructure import secure_files
from fire_safety_backend.infrastructure.generators.redline_docx import (
    AUTHOR,
    MIN_QUOTE_LEN,
    _usable_findings,
    build_redline_docx,
)

WHEN = "2026-08-07T12:00:00Z"

QUOTE = "За просрочку сдачи работ Подрядчик уплачивает пеню 1% от цены Договора за каждый день просрочки"
FIX = "Пеня 0,1% от стоимости несданного этапа за каждый день просрочки, но не более 10% от цены Договора"

TABLE_QUOTE = (
    "Гарантийный срок на смонтированные системы составляет 12 месяцев с даты подписания акта"
)
TABLE_FIX = "Гарантийный срок на смонтированные системы составляет 24 месяца с даты подписания акта"


@pytest.fixture(autouse=True)
def _isolated_output(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "out")
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(config, "WORK_DIR", tmp_path / "work")
    config.WORK_DIR.mkdir(parents=True, exist_ok=True)


def _source(tmp_path: Path) -> Path:
    """Договор, в котором цитата НАМЕРЕННО разрезана между run'ами.

    Word дробит абзац по своей логике — из-за смены шрифта, проверки
    орфографии, предыдущих правок, — и цитата почти всегда оказывается
    разорванной. Ровно этот случай и должен работать.
    """
    path = tmp_path / "Договор подряда.docx"
    doc = Document()
    doc.add_heading("ДОГОВОР ПОДРЯДА", level=1)
    doc.add_paragraph("1. Предмет договора. Подрядчик выполняет монтаж систем пожаротушения.")

    para = doc.add_paragraph("5.1. ")
    para.add_run(QUOTE[:30]).bold = True
    para.add_run(QUOTE[30:])
    para.add_run(". Настоящий пункт применяется с даты подписания.")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "7.2"
    table.rows[0].cells[1].text = TABLE_QUOTE + "."

    doc.save(path)
    return path


def _findings(*pairs: tuple[str, str], found: bool = True) -> list[dict]:
    return [
        {
            "критичность": "красный",
            "цитата_из_договора": q,
            "предложение_правки": f,
            "_цитата_найдена": found,
        }
        for q, f in pairs
    ]


def _open(logical: Path) -> Document:
    return Document(io.BytesIO(secure_files.load(logical)))


def _visible(doc: Document, *, accept: bool) -> str:
    """Текст документа глазами того, кто принял или отклонил все правки.

    Принял — видны вставки (w:t внутри w:ins) и неизменный текст, удалённого
    нет. Отклонил — наоборот: виден w:delText, вставок нет.
    """
    out: list[str] = []
    for node in doc.element.iter():
        tag = node.tag.split("}")[-1]
        in_ins = any(a.tag.split("}")[-1] == "ins" for a in node.iterancestors())
        shown = (tag == "t" and (accept or not in_ins)) or (tag == "delText" and not accept)
        if shown:
            out.append(node.text or "")
    return "".join(out)


# --- главное: два состояния документа ---


def test_accepting_changes_gives_new_wording(tmp_path: Path) -> None:
    out, applied, total = build_redline_docx(_findings((QUOTE, FIX)), _source(tmp_path), WHEN)
    assert (applied, total) == (1, 1)
    text = _visible(_open(out), accept=True)
    assert FIX in text
    assert QUOTE not in text


def test_rejecting_changes_gives_original(tmp_path: Path) -> None:
    """Отклонивший обязан получить ИСХОДНЫЙ договор без следов правки."""
    out, _, _ = build_redline_docx(_findings((QUOTE, FIX)), _source(tmp_path), WHEN)
    text = _visible(_open(out), accept=False)
    assert QUOTE in text
    assert FIX not in text


def test_untouched_text_survives_both_ways(tmp_path: Path) -> None:
    """Голова и хвост абзаца не должны пострадать при разрезании run'ов."""
    out, _, _ = build_redline_docx(_findings((QUOTE, FIX)), _source(tmp_path), WHEN)
    doc = _open(out)
    for accept in (True, False):
        text = _visible(doc, accept=accept)
        assert "5.1. " in text
        assert "Настоящий пункт применяется с даты подписания." in text
        assert "Подрядчик выполняет монтаж систем пожаротушения." in text


def test_redline_works_inside_tables(tmp_path: Path) -> None:
    """В договорах сроки и гарантии живут в таблицах приложений."""
    out, applied, _ = build_redline_docx(
        _findings((TABLE_QUOTE, TABLE_FIX)), _source(tmp_path), WHEN
    )
    assert applied == 1
    assert TABLE_FIX in _visible(_open(out), accept=True)


def test_several_findings_all_applied(tmp_path: Path) -> None:
    out, applied, total = build_redline_docx(
        _findings((QUOTE, FIX), (TABLE_QUOTE, TABLE_FIX)), _source(tmp_path), WHEN
    )
    assert (applied, total) == (2, 2)
    text = _visible(_open(out), accept=True)
    assert FIX in text
    assert TABLE_FIX in text


# --- структура разметки: этого не видно по тексту ---


def test_deleted_text_uses_deltext_element(tmp_path: Path) -> None:
    """Внутри w:del обязан быть w:delText.

    С w:t Word считает файл повреждённым и открывает его БЕЗ правок — то есть
    юрист отправит контрагенту исходный договор, думая, что отправил редлайн.
    """
    out, _, _ = build_redline_docx(_findings((QUOTE, FIX)), _source(tmp_path), WHEN)
    xml = _open(out).element.xml
    assert "<w:delText" in xml
    for chunk in xml.split("<w:del ")[1:]:
        body = chunk.split("</w:del>")[0]
        assert "<w:t>" not in body and "<w:t " not in body


def test_change_ids_are_unique(tmp_path: Path) -> None:
    """Повторный w:id склеивает разные правки в одну — принять их по
    отдельности станет нельзя."""
    out, _, _ = build_redline_docx(
        _findings((QUOTE, FIX), (TABLE_QUOTE, TABLE_FIX)), _source(tmp_path), WHEN
    )
    xml = _open(out).element.xml
    ids = [part.split('"')[0] for part in xml.split('w:id="')[1:]]
    marks = [i for i in ids if i.isdigit()]
    assert len(marks) == len(set(marks)), f"повторяются идентификаторы правок: {marks}"


def test_author_and_date_are_set(tmp_path: Path) -> None:
    out, _, _ = build_redline_docx(_findings((QUOTE, FIX)), _source(tmp_path), WHEN)
    xml = _open(out).element.xml
    assert f'w:author="{AUTHOR}"' in xml
    assert f'w:date="{WHEN}"' in xml


# --- что ОБЯЗАНО быть пропущено ---


def test_unverified_quote_is_skipped(tmp_path: Path) -> None:
    """Цитата не найдена в оригинале дословно — модель её перефразировала.

    Подставить правку не в то место договора хуже, чем не подставить.
    """
    out, applied, total = build_redline_docx(
        _findings((QUOTE, FIX), found=False), _source(tmp_path), WHEN
    )
    assert (applied, total) == (0, 0)
    assert QUOTE in _visible(_open(out), accept=True)


@pytest.mark.parametrize(
    ("quote", "fix", "why"),
    [
        (QUOTE, "", "находка есть, формулировки нет"),
        (QUOTE, QUOTE, "замена совпадает с цитатой — пустая пометка в Word"),
        ("Подрядчик обязан", "Подрядчик вправе", "цитата короче предела, уедет не в тот пункт"),
    ],
)
def test_unusable_findings_are_skipped(quote: str, fix: str, why: str) -> None:
    assert _usable_findings(_findings((quote, fix))) == [], why


def test_longer_quote_applied_first() -> None:
    """Если одна цитата содержится в другой, короткая не должна разрезать
    длинную раньше времени."""
    short = "а" * (MIN_QUOTE_LEN + 5)
    long = short + " и ещё продолжение этого же пункта договора"
    pairs = _usable_findings(_findings((short, "б"), (long, "в")))
    assert pairs[0][0] == long


def test_quote_absent_from_docx_is_counted_but_not_applied(tmp_path: Path) -> None:
    """Цитата была в извлечённом тексте, но в самом DOCX не нашлась.

    Так бывает, когда она пересекает границу абзаца. Интерфейс обязан
    показать разницу: эти правки юрист внесёт руками.
    """
    missing = "Такого предложения в договоре нет совершенно точно, ни одного слова"
    out, applied, total = build_redline_docx(
        _findings((missing, "замена")), _source(tmp_path), WHEN
    )
    assert (applied, total) == (0, 1)
    assert out.exists()


def test_pdf_source_is_refused(tmp_path: Path) -> None:
    """Из PDF редлайн собрать нельзя: править нечего.

    Честнее отказать, чем отдать пересобранную простыню, выдав её за договор.
    """
    pdf = tmp_path / "Договор.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    with pytest.raises(ValueError, match="DOCX"):
        build_redline_docx(_findings((QUOTE, FIX)), pdf, WHEN)
