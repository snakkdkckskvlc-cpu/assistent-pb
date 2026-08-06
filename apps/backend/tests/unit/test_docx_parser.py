"""Чтение DOCX не должно молча терять текст.

Договор от контрагента С ПРАВКАМИ рецензирования — норма при согласовании, а не
редкость. `python-docx` собирает `paragraph.text` только из прямых потомков
абзаца, а вставленный текст лежит внутри `w:ins` и туда не попадает; удалённый
лежит в `w:delText`, а не в `w:t`.

Проверено экспериментально до починки: в «в течение 30 банковских дней» число
ИСЧЕЗАЛО, оставался двойной пробел. Модель получала текст с дырами и не знала,
что они есть, — ответ при этом выглядел полным. Это худший вид отказа.
"""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from fire_safety_backend.infrastructure.parsers.docx_parser import extract_docx  # noqa: E402


def _run_with_text(paragraph, tag: str, text: str):
    """Прогон с текстом в заданном теге (w:t — обычный, w:delText — удалённый)."""
    run = paragraph._p.makeelement(qn("w:r"), {})
    node = paragraph._p.makeelement(qn(tag), {})
    node.text = text
    run.append(node)
    return run


def _wrap(paragraph, tag: str, run):
    """Оборачивает прогон в w:ins или w:del — так Word помечает правку."""
    wrapper = paragraph._p.makeelement(
        qn(tag),
        {qn("w:id"): "1", qn("w:author"): "Заказчик", qn("w:date"): "2026-01-01T00:00:00Z"},
    )
    wrapper.append(run)
    paragraph._p.append(wrapper)


def test_inserted_text_is_not_lost(tmp_path: Path) -> None:
    """Вставленный при рецензировании текст лежит внутри w:ins."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("в течение ")
    _wrap(p, "w:ins", _run_with_text(p, "w:t", "30 банковских"))
    p.add_run(" дней")
    path = tmp_path / "с_правками.docx"
    doc.save(path)

    assert extract_docx(path) == "в течение 30 банковских дней"


def test_deleted_text_is_kept(tmp_path: Path) -> None:
    """Зачёркнутый в согласовании пункт для разбора рисков не менее важен, чем
    оставшийся, — а иногда важнее: именно из него видно, что сторона пыталась
    убрать неустойку."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Подрядчик уплачивает ")
    _wrap(p, "w:del", _run_with_text(p, "w:delText", "неустойку 1 %"))
    path = tmp_path / "с_удалением.docx"
    doc.save(path)

    assert "неустойку 1 %" in extract_docx(path)


def test_table_stays_next_to_its_paragraph(tmp_path: Path) -> None:
    """Раньше абзацы и таблицы выгружались раздельно, и таблица уезжала в конец,
    отрываясь от пункта, который на неё ссылается. Для разбора договора связь
    пункта с таблицей и есть содержание."""
    doc = Document()
    doc.add_paragraph("Порядок оплаты приведён в таблице.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Этап"
    table.rows[0].cells[1].text = "Срок"
    doc.add_paragraph("Последний абзац.")
    path = tmp_path / "с_таблицей.docx"
    doc.save(path)

    lines = extract_docx(path).splitlines()
    assert lines[0] == "Порядок оплаты приведён в таблице."
    assert lines[1] == "Этап | Срок"
    assert lines[2] == "Последний абзац."


def test_plain_document_reads_as_before(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("")
    doc.add_paragraph("Второй абзац.")
    path = tmp_path / "обычный.docx"
    doc.save(path)

    assert extract_docx(path) == "Первый абзац.\nВторой абзац."
