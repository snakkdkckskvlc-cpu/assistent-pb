"""Правка DOCX на месте: исправляем ошибки, не трогая оформление.

Регрессия на реальную жалобу: проверка орфографии отдавала «исправленный
текст» простынёй, и пользователь должен был сам переносить правки в свой
документ, теряя шрифты, отступы и таблицы.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from fire_safety_backend.services.docx_edit import apply_corrections_to_docx


@pytest.fixture
def make_docx(tmp_path):
    def _make(build) -> object:
        doc = Document()
        build(doc)
        path = tmp_path / "src.docx"
        doc.save(str(path))
        return path

    return _make


def test_correction_preserves_bold_italic_size_and_colour(make_docx, tmp_path):
    def build(doc):
        p = doc.add_paragraph()
        r1 = p.add_run("Работы на ")
        r1.bold = True
        r1.font.size = Pt(14)
        r2 = p.add_run("обьекте ")
        r2.italic = True
        r2.font.color.rgb = RGBColor(0xC0, 0, 0)
        r3 = p.add_run("завершены.")
        r3.underline = True

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "обьекте", "after": "объекте"}]
    )
    assert report.applied == 1

    runs = Document(str(out)).paragraphs[0].runs
    assert "объекте" in runs[1].text
    assert runs[0].bold is True and runs[0].font.size == Pt(14)
    assert runs[1].italic is True and runs[1].font.color.rgb == RGBColor(0xC0, 0, 0)
    assert runs[2].underline is True


def test_word_split_across_runs_is_fixed(make_docx, tmp_path):
    """Word произвольно дробит абзац на прогоны — слово с ошибкой запросто
    оказывается разрезанным пополам, и наивный поиск по прогону его не найдёт."""

    def build(doc):
        p = doc.add_paragraph()
        p.add_run("Работы на об")
        p.add_run("ьекте")
        p.add_run(" завершены")

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "обьекте", "after": "объекте"}]
    )
    assert report.applied == 1
    assert report.split_formatting == 1, "правка на границе прогонов должна быть отмечена"
    assert Document(str(out)).paragraphs[0].text == "Работы на объекте завершены"


def test_all_occurrences_are_fixed_not_just_the_first(make_docx, tmp_path):
    """Одна опечатка обычно повторяется по документу. Исправить только первое
    вхождение значит отдать наполовину вычитанный файл."""

    def build(doc):
        doc.add_paragraph("Срок истекает в течении месяца.")
        doc.add_paragraph("Оплата в течении 30 дней и в течении квартала.")

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "в течении", "after": "в течение"}]
    )
    assert report.applied == 3
    assert "в течении" not in "\n".join(p.text for p in Document(str(out)).paragraphs)


def test_replacement_containing_the_original_does_not_loop(make_docx, tmp_path):
    """«месяца» → «месяца работ» содержит исходный фрагмент: повторный проход
    нашёл бы собственный результат и зациклился."""

    def build(doc):
        doc.add_paragraph("Срок — один месяца.")

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "месяца", "after": "месяца работ"}]
    )
    assert report.applied == 1
    assert Document(str(out)).paragraphs[0].text == "Срок — один месяца работ."


def test_corrections_in_tables_and_headers(make_docx, tmp_path):
    """Реквизиты компании в бланке лежат в таблице и колонтитуле — опечатка
    там уедет к контрагенту так же, как из тела письма."""

    def build(doc):
        doc.add_paragraph("Тело письма.")
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "Адрес обьекта уточняется."
        doc.sections[0].header.paragraphs[0].text = "ООО «ПожСервис», обьект №1"

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "обьект", "after": "объект"}]
    )
    assert report.applied == 2
    result = Document(str(out))
    assert "объекта" in result.tables[0].cell(0, 0).text
    assert "объект" in result.sections[0].header.paragraphs[0].text


def test_missing_fragment_is_reported_not_silently_dropped(make_docx, tmp_path):
    def build(doc):
        doc.add_paragraph("Текст без ошибок.")

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build), out, [{"before": "которого нет", "after": "x"}]
    )
    assert report.applied == 0
    assert report.not_found == ["которого нет"]


def test_original_file_is_never_modified(make_docx, tmp_path):
    def build(doc):
        doc.add_paragraph("Работы на обьекте.")

    src = make_docx(build)
    before = src.read_bytes()
    apply_corrections_to_docx(
        src, tmp_path / "out.docx", [{"before": "обьекте", "after": "объекте"}]
    )
    assert src.read_bytes() == before, "оригинал должен остаться нетронутым"


def test_noop_corrections_are_skipped(make_docx, tmp_path):
    def build(doc):
        doc.add_paragraph("Текст.")

    out = tmp_path / "out.docx"
    report = apply_corrections_to_docx(
        make_docx(build),
        out,
        [{"before": "Текст", "after": "Текст"}, {"before": "", "after": "x"}],
    )
    assert report.applied == 0
    assert report.not_found == []
