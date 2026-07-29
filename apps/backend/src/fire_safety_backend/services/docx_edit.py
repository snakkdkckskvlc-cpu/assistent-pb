"""Правка DOCX на месте: заменить фрагменты, не тронув оформление.

Зачем. Проверка орфографии до сих пор отдавала «исправленный текст» простыней:
пользователь получал голый текст и должен был сам переносить правки в свой
документ, теряя шрифты, отступы, таблицы и нумерацию. Здесь правки
применяются к КОПИИ исходного файла — меняются только те символы, где была
ошибка.

Как сохраняется оформление. В DOCX формат хранится не у слова, а у «прогона»
(run) — куска текста с одинаковым набором свойств. Word дробит абзац на
прогоны произвольно: одно слово может оказаться разрезанным на три прогона
из-за подчёркивания проверки правописания или следа от правки. Поэтому
поиск идёт по СКЛЕЕННОМУ тексту абзаца, а замена — по карте «позиция в
абзаце → прогон», и новый текст кладётся в первый задетый прогон, чьи
свойства и наследует. Остальные задетые прогоны затираются.

Ограничение, важное для понимания: если ошибка пришлась на границу прогонов
с РАЗНЫМ оформлением (половина слова жирная), всё исправление получит формат
первого из них. Это лучше, чем отказ от правки, но не идеально; такие случаи
считаются и возвращаются в отчёте.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from pathlib import Path

log = logging.getLogger(__name__)

# Предохранитель от бесконечного цикла: если after СОДЕРЖИТ before (например
# «в течении» → «в течении месяца»), поиск найдёт замену снова и снова.
_MAX_REPLACEMENTS_PER_CORRECTION = 200


@dataclass
class EditReport:
    """Что удалось применить, а что нет."""

    applied: int = 0
    not_found: list[str] = None
    ambiguous: list[str] = None
    split_formatting: int = 0

    def __post_init__(self) -> None:
        if self.not_found is None:
            self.not_found = []
        if self.ambiguous is None:
            self.ambiguous = []

    def as_dict(self) -> dict:
        return {
            "применено": self.applied,
            "не_найдено": self.not_found,
            "неоднозначно": self.ambiguous,
            "правок_на_границе_форматов": self.split_formatting,
        }


def _normalize(text: str) -> str:
    """Пробелы к одному виду. Word щедро сыплет неразрывными пробелами и
    мягкими переносами, из-за чего дословный поиск фрагмента промахивается."""
    return re.sub(r"[\s ​­]+", " ", text)


def _paragraph_runs(paragraph):
    """Прогоны абзаца вместе с диапазоном, который каждый занимает в его тексте."""
    spans = []
    pos = 0
    for run in paragraph.runs:
        length = len(run.text)
        spans.append((pos, pos + length, run))
        pos += length
    return spans


def _replace_in_paragraph(paragraph, before: str, after: str, report: EditReport) -> bool:
    """Заменяет ПЕРВОЕ вхождение before на after внутри одного абзаца."""
    spans = _paragraph_runs(paragraph)
    if not spans:
        return False
    full = "".join(run.text for _, _, run in spans)

    start = full.find(before)
    if start < 0:
        # Второй заход по нормализованным пробелам: фрагмент от модели почти
        # всегда приходит с обычными пробелами, а в документе может стоять
        # неразрывный.
        pattern = r"[\s ​­]+".join(re.escape(w) for w in before.split())
        match = re.search(pattern, full)
        if not match:
            return False
        start, end = match.span()
    else:
        end = start + len(before)

    touched = [(s, e, r) for s, e, r in spans if s < end and e > start]
    if not touched:
        return False
    if len({id(r) for _, _, r in touched}) > 1:
        report.split_formatting += 1

    first_start, _, first_run = touched[0]
    # Хвосты прогона, не попавшие в заменяемый фрагмент, обязаны уцелеть:
    # ошибка редко занимает прогон целиком.
    head = first_run.text[: start - first_start]
    last_start, last_end, last_run = touched[-1]
    tail = last_run.text[end - last_start :] if end > last_start else ""

    first_run.text = head + after + (tail if last_run is first_run else "")
    for _, _, run in touched[1:]:
        run.text = "" if run is not last_run else tail
    return True


def _iter_paragraphs(doc):
    """Все абзацы документа: тело, таблицы, колонтитулы.

    Колонтитулы включены намеренно — в бланках компании там реквизиты, и
    опечатка в них так же попадёт к контрагенту.
    """
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def apply_corrections_to_docx(source: Path, output: Path, corrections: list[dict]) -> EditReport:
    """Копирует DOCX и применяет к копии правки вида {before, after}.

    Оригинал не трогается никогда: пользователь должен иметь возможность
    сравнить и откатиться.
    """
    from docx import Document

    report = EditReport()
    doc = Document(str(source))
    paragraphs = list(_iter_paragraphs(doc))

    for correction in corrections:
        before = str(correction.get("before", "") or "")
        after = str(correction.get("after", "") or "")
        if not before or before == after:
            continue

        # Правка применяется ко ВСЕМ вхождениям, а не к первому. Одна и та же
        # опечатка обычно повторяется по всему документу («в течении» в двух
        # абзацах — две ошибки, а не одна), и исправлять только первое значило
        # бы отдать пользователю наполовину вычитанный файл. Обратный риск —
        # что фрагмент где-то окажется законным — мал: `before` приходит от
        # LanguageTool и модели как конкретное место ошибки, а не как слово
        # общего употребления.
        matches = [p for p in paragraphs if before in "".join(r.text for r in p.runs)]
        if not matches:
            norm_before = _normalize(before)
            matches = [p for p in paragraphs if norm_before in _normalize(p.text)]
        if not matches:
            report.not_found.append(before)
            continue

        # Если исправление содержит в себе исходный фрагмент («в течении» →
        # «в течении месяца»), повторный проход нашёл бы собственный результат
        # и зациклился. В таком случае — строго одна замена на абзац.
        repeatable = before not in after
        touched = 0
        for paragraph in matches:
            replaced = _replace_in_paragraph(paragraph, before, after, report)
            touched += int(replaced)
            while repeatable and replaced and touched < _MAX_REPLACEMENTS_PER_CORRECTION:
                replaced = _replace_in_paragraph(paragraph, before, after, report)
                touched += int(replaced)
        if touched:
            report.applied += touched
        else:
            report.not_found.append(before)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    log.info("Правки в DOCX: %s", report.as_dict())
    return report
