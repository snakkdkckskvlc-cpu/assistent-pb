"""Печать путевого листа: типовые межотраслевые формы № 3 и 4-С.

Бланк собирается кодом, а не подстановкой в файл-образец. Причина техническая:
действующие бланки компании лежат в старых двоичных форматах — легковой в .doc,
грузовой в .xls, — а ни python-docx, ни openpyxl их не открывают. Конвертировать
их в .docx/.xlsx нечем: ни LibreOffice, ни Word в окружении приложения нет, а
на рабочих местах конвертация вручную означала бы, что шаблон у каждого свой.

Поэтому формы воспроизведены разметкой: тот же состав и тот же порядок полей,
те же формулировки подписей, что в бумажном бланке (формы утверждены
постановлением Госкомстата России № 78 от 28.11.1997, ОКУД 0345001 и 0345004).
Это НЕ побайтовая копия файла компании: расположение блоков на странице может
отличаться. Если бланк потребуется один в один, правильный путь — сохранить
исходники как .docx/.xlsx с плейсхолдерами и перейти на подстановку, как
сделано для писем (letter_docx.py).

Подписи не печатаются — под них оставлены линейки. Расшифровки подписей
печатаются: это данные, которые ведёт приложение.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

if TYPE_CHECKING:
    from pathlib import Path

_FONT = "Times New Roman"
_SIZE = Pt(8)
_SIZE_SMALL = Pt(6.5)
_SIZE_TITLE = Pt(11)

# Линейка под подпись. Символ подчёркивания, а не пустая ячейка с рамкой:
# бланк печатают и заполняют ручкой, и человеку нужна видимая строка.
_RULE = "_" * 22


def _fmt_date(value: Any) -> str:
    """ISO-дата в вид бланка. Пусто и мусор одинаково дают прочерк."""
    if not value:
        return ""
    text = str(value)[:10]
    parts = text.split("-")
    if len(parts) == 3 and all(p.isdigit() for p in parts):
        return f"{parts[2]}.{parts[1]}.{parts[0]}"
    return text


def _fmt_dt(value: Any) -> str:
    if not value:
        return ""
    text = str(value).replace("T", " ")
    return f"{_fmt_date(text)} {text[11:16]}".strip()


def _fmt_time(value: Any) -> str:
    if not value:
        return ""
    return str(value).replace("T", " ")[11:16]


def _fmt_num(value: Any, digits: int = 2) -> str:
    """Число в русском виде: разделитель дробной части — запятая."""
    if value is None or value == "":
        return ""
    try:
        text = f"{float(value):.{digits}f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return str(value)
    return text.replace(".", ",") or "0"


def _fmt_minutes(value: Any) -> str:
    """Минуты в «ч. мин» — так подписана графа бланка."""
    if value is None or value == "":
        return ""
    try:
        total = int(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{total // 60} ч {total % 60:02d} мин"


def _style_cell(cell, *, bold: bool = False, size=_SIZE, align=None) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = _FONT
            run.font.size = size
            run.font.bold = bold
            # У кириллицы в DOCX своя настройка шрифта: без w:cs/w:eastAsia
            # Word подставляет свой и текст в таблицах «плывёт».
            run._element.rPr.rFonts.set(qn("w:cs"), _FONT)


def _para(doc, text: str, *, bold=False, size=_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT, space=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space)
    run = p.add_run(text)
    run.font.name = _FONT
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:cs"), _FONT)
    return p


def _grid(doc, rows: list[list[str]], widths: list[float], *, header: bool = False):
    """Таблица с рамками. rows[0] — шапка, если header=True."""
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r, line in enumerate(rows):
        for c, value in enumerate(line):
            cell = table.cell(r, c)
            cell.width = Cm(widths[c])
            cell.text = "" if value is None else str(value)
            _style_cell(
                cell,
                bold=header and r == 0,
                size=_SIZE_SMALL if header and r == 0 else _SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER if header and r == 0 else None,
            )
    return table


def _pairs(doc, items: list[tuple[str, str]], *, cols: int = 2):
    """Блок «подпись — значение». Это основной способ показать поля бланка,
    у которого в бумаге нет табличной сетки: слева наименование графы, справа
    то, что в неё вписано."""
    items = [(label, value) for label, value in items]
    if not items:
        return
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols * 2)
    table.style = "Table Grid"
    table.autofit = False
    width_label, width_value = (6.4, 7.0) if cols == 2 else (8.0, 18.0)
    for index, (label, value) in enumerate(items):
        r, c = index % rows, (index // rows) * 2
        label_cell = table.cell(r, c)
        label_cell.width = Cm(width_label)
        label_cell.text = label
        _style_cell(label_cell, size=_SIZE_SMALL)
        value_cell = table.cell(r, c + 1)
        value_cell.width = Cm(width_value)
        value_cell.text = value or ""
        _style_cell(value_cell)
    # Хвост, если полей нечётное число: пустые ячейки без подписей.
    for index in range(len(items), rows * cols):
        r, c = index % rows, (index // rows) * 2
        _style_cell(table.cell(r, c), size=_SIZE_SMALL)
        _style_cell(table.cell(r, c + 1))
    return table


def _signature_line(doc, caption: str, name: str) -> None:
    """Линейка под подпись с расшифровкой. Саму подпись ставят на бумаге."""
    _para(doc, f"{caption} {_RULE}  {name or _RULE}")
    _para(
        doc,
        " " * len(caption) + "         подпись            расшифровка подписи",
        size=_SIZE_SMALL,
        space=4,
    )


def _landscape(doc) -> None:
    """Обе формы печатаются альбомно — так же, как настроены бланки компании."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.0)
    section.top_margin = section.bottom_margin = Cm(1.0)


def _header(doc, wb: dict, *, form_title: str, form_number: str, okud: str) -> None:
    org_line = " ".join(
        x for x in (wb.get("org_name", ""), wb.get("org_address", ""), wb.get("org_phone", "")) if x
    )
    _para(doc, "Место для штампа организации", size=_SIZE_SMALL)
    _para(
        doc,
        f"Типовая межотраслевая форма № {form_number}. "
        "Утверждена постановлением Госкомстата России от 28.11.1997 № 78",
        size=_SIZE_SMALL,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    series = wb.get("series") or "______"
    number = wb.get("number") or "______"
    _para(
        doc,
        f"{form_title} {series} № {number}",
        bold=True,
        size=_SIZE_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space=2,
    )
    period = _fmt_date(wb.get("date_from"))
    if wb.get("date_to") and wb.get("date_to") != wb.get("date_from"):
        period = f"{period} по {_fmt_date(wb.get('date_to'))}"
        _para(doc, f"Срок действия с {period}", align=WD_ALIGN_PARAGRAPH.CENTER, space=4)
    else:
        _para(doc, f"Срок действия {period}", align=WD_ALIGN_PARAGRAPH.CENTER, space=4)

    _pairs(
        doc,
        [
            ("Организация (наименование, адрес, телефон)", org_line),
            ("Форма по ОКУД", okud),
            ("по ОКПО", wb.get("org_okpo", "")),
            ("ОГРН", wb.get("org_ogrn", "")),
        ],
    )
    _para(doc, "", space=4)


def _vehicle_driver_block(wb: dict, *, cargo: bool) -> list[tuple[str, str]]:
    items = [
        ("Марка автомобиля", wb.get("vehicle_mark", "")),
        ("Государственный номерной знак", wb.get("vehicle_plate", "")),
        ("Гаражный номер", wb.get("garage_number", "")),
        ("Водитель (фамилия, имя, отчество)", wb.get("driver_name", "")),
        ("Табельный номер", wb.get("driver_tab_number", "")),
        (
            "Удостоверение №",
            " ".join(
                x
                for x in (wb.get("driver_licence_series", ""), wb.get("driver_licence_number", ""))
                if x
            ),
        ),
        ("выдано", _fmt_date(wb.get("driver_licence_issued_at"))),
        ("Класс", wb.get("driver_licence_class", "")),
        ("СНИЛС", wb.get("driver_snils", "")),
        ("Лицензионная карточка", wb.get("driver_licence_card", "")),
    ]
    if cargo:
        items += [
            ("Колонна", wb.get("column_number", "")),
            ("Бригада", wb.get("brigade", "")),
            ("Режим работы", wb.get("work_mode", "")),
            ("Сопровождающие лица", wb.get("escorts", "")),
        ]
    return items


def _fuel_rows(wb: dict, *, cargo: bool) -> list[tuple[str, str]]:
    rows = [
        ("Горючее: марка", wb.get("fuel_brand", "")),
        ("код", wb.get("fuel_code", "")),
        ("Выдано, л", _fmt_num(wb.get("fuel_issued_l"))),
        ("Остаток при выезде, л", _fmt_num(wb.get("fuel_start_l"))),
        ("Остаток при возвращении, л", _fmt_num(wb.get("fuel_end_l"))),
    ]
    if cargo:
        rows += [
            ("Сдано, л", _fmt_num(wb.get("fuel_returned_l"))),
            (
                "Коэффициент изменения нормы: спецоборудования",
                _fmt_num(wb.get("fuel_coeff_equipment")),
            ),
            ("двигателя", _fmt_num(wb.get("fuel_coeff_engine"))),
            ("Время работы спецоборудования", _fmt_minutes(wb.get("equipment_time_min"))),
            ("Время работы двигателя", _fmt_minutes(wb.get("engine_time_min"))),
        ]
    else:
        rows += [("По заправочному листу №", wb.get("fuel_sheet_number", ""))]

    # Вписанная человеком цифра главнее посчитанной: если расход по норме
    # уже согласован с бухгалтерией, бланк обязан показать именно её.
    norm = wb.get("fuel_used_norm_l")
    if norm in (None, ""):
        norm = wb.get("fuel_by_norm_l")
    fact = wb.get("fuel_used_fact_l")
    if fact in (None, ""):
        fact = wb.get("fuel_balance_l")
    saving = wb.get("fuel_saving_l")
    rows += [
        ("Расход: по норме, л", _fmt_num(norm)),
        ("фактически, л", _fmt_num(fact)),
        # Экономия и перерасход — две РАЗНЫЕ графы бланка, и в одну из них
        # ставится прочерк. Печатать одно число со знаком было бы неверно:
        # бухгалтерия читает именно эти две графы.
        ("Экономия, л", _fmt_num(saving) if saving is not None and saving > 0 else ""),
        (
            "Перерасход, л",
            _fmt_num(abs(saving)) if saving is not None and saving < 0 else "",
        ),
    ]
    return rows


def _medical_block(doc, wb: dict) -> None:
    _para(
        doc,
        "Предрейсовый и послерейсовый медицинский осмотр, контроль технического состояния",
        bold=True,
        space=2,
    )
    _pairs(
        doc,
        [
            ("Предрейсовый медосмотр: отметка о прохождении", wb.get("medical_pre_mark", "")),
            ("дата, время", _fmt_dt(wb.get("medical_pre_at"))),
            ("должность медработника", wb.get("medical_pre_position", "")),
            ("расшифровка подписи", wb.get("medical_pre_name", "")),
            ("Послерейсовый медосмотр: отметка о прохождении", wb.get("medical_post_mark", "")),
            ("дата, время", _fmt_dt(wb.get("medical_post_at"))),
            ("должность медработника", wb.get("medical_post_position", "")),
            ("расшифровка подписи", wb.get("medical_post_name", "")),
            ("Предрейсовый контроль технического состояния", wb.get("tech_control_mark", "")),
            ("дата, время", _fmt_dt(wb.get("tech_control_at"))),
            ("Контролёр технического состояния", wb.get("tech_control_name", "")),
            ("При возвращении автомобиль", wb.get("return_condition", "")),
        ],
    )
    _para(doc, "", space=4)
    _signature_line(doc, "Выезд разрешен. Механик", wb.get("mechanic_name", ""))
    _signature_line(
        doc,
        "Автомобиль в технически исправном состоянии принял. Водитель",
        wb.get("accepted_by_driver", ""),
    )
    _signature_line(doc, "Диспетчер-нарядчик", wb.get("dispatcher_name", ""))
    _signature_line(doc, "Автомобиль сдал. Водитель", wb.get("vehicle_handed_by", ""))
    _signature_line(doc, "Автомобиль принял. Механик", wb.get("vehicle_taken_by", ""))


def _build_form3(doc, wb: dict, trips: list[dict]) -> None:
    _header(
        doc,
        wb,
        form_title="ПУТЕВОЙ ЛИСТ ЛЕГКОВОГО АВТОМОБИЛЯ",
        form_number="3",
        okud="0345001",
    )
    _pairs(doc, _vehicle_driver_block(wb, cargo=False))
    _para(doc, "", space=4)

    _para(doc, "Задание водителю", bold=True, space=2)
    _pairs(
        doc,
        [
            ("Вид сообщения", wb.get("communication_type", "")),
            ("Вид перевозки", wb.get("transport_type", "")),
            ("В распоряжение (наименование)", wb.get("customer_name", "")),
            ("Адрес подачи", wb.get("pickup_address", "")),
            ("Дата, время выезда из гаража", _fmt_dt(wb.get("departure_at"))),
            ("Дата, время возвращения в гараж", _fmt_dt(wb.get("return_at"))),
            ("Показания одометра при выезде, км", _fmt_num(wb.get("odometer_start_km"), 1)),
            ("Показания одометра при возвращении, км", _fmt_num(wb.get("odometer_end_km"), 1)),
        ],
    )
    _para(doc, "", space=4)

    _para(doc, "Движение горючего", bold=True, space=2)
    _pairs(doc, _fuel_rows(wb, cargo=False))
    _para(doc, "", space=4)

    _medical_block(doc, wb)

    if wb.get("delay_notes"):
        _para(doc, "", space=4)
        _para(
            doc, "Опоздания, ожидания, простои, заезды в гараж и прочие отметки", bold=True, space=2
        )
        _para(doc, wb["delay_notes"])

    doc.add_page_break()
    _para(
        doc,
        "Оборотная сторона формы № 3",
        size=_SIZE_SMALL,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space=4,
    )
    rows = [
        [
            "Номер по порядку",
            "Код заказчика",
            "Место отправления",
            "Место назначения",
            "Время выезда",
            "Время возвращения",
            "Пройдено, км",
            "Подпись лица, пользовавшегося автомобилем",
        ]
    ]
    for index, trip in enumerate(trips, start=1):
        rows.append(
            [
                str(trip.get("seq") or index),
                trip.get("customer_code", ""),
                trip.get("place_from_name") or "",
                trip.get("place_to_name") or trip.get("destination_text") or "",
                _fmt_time(trip.get("departed_at")),
                _fmt_time(trip.get("returned_at")),
                _fmt_num(trip.get("distance_km"), 1),
                trip.get("user_name", ""),
            ]
        )
    # Пустые строки, чтобы в напечатанный бланк можно было дописать ручкой.
    for _ in range(max(0, 8 - len(trips))):
        rows.append([""] * 8)
    _grid(doc, rows, [1.6, 1.8, 4.6, 4.6, 2.2, 2.4, 2.0, 4.4], header=True)

    _para(doc, "", space=4)
    _para(doc, "Результат работы автомобиля за смену", bold=True, space=2)
    _pairs(
        doc,
        [
            ("Всего в наряде", _fmt_minutes(wb.get("total_time_min"))),
            ("Пройдено, км", _fmt_num(wb.get("run_total_km"), 1)),
        ],
    )
    _para(doc, "", space=4)
    _para(doc, "Расчёт заработной платы", bold=True, space=2)
    _pairs(
        doc,
        [
            ("За километраж, руб. коп.", _fmt_num(wb.get("salary_km_rub"))),
            ("За часы, руб. коп.", _fmt_num(wb.get("salary_hours_rub"))),
            ("Итого, руб. коп.", _fmt_num(wb.get("salary_total_rub"))),
            ("Должность", wb.get("calc_position", "")),
        ],
    )
    _para(doc, "", space=4)
    _signature_line(doc, "Расчёт произвёл", wb.get("calc_name", ""))


def _build_form4s(doc, wb: dict, trips: list[dict]) -> None:
    _header(
        doc,
        wb,
        form_title="ПУТЕВОЙ ЛИСТ ГРУЗОВОГО АВТОМОБИЛЯ",
        form_number="4-С",
        okud="0345004",
    )
    _pairs(doc, _vehicle_driver_block(wb, cargo=True))
    _para(doc, "", space=4)

    trailers = wb.get("trailers") or []
    if trailers:
        _para(doc, "Прицепы", bold=True, space=2)
        rows = [["№", "Марка", "Регистрационный №", "Серия", "Код марки"]]
        for index, trailer in enumerate(trailers, start=1):
            rows.append(
                [
                    str(index),
                    trailer.get("mark", ""),
                    trailer.get("reg_number", ""),
                    trailer.get("series", ""),
                    trailer.get("code", ""),
                ]
            )
        _grid(doc, rows, [1.2, 6.0, 5.0, 3.0, 3.0], header=True)
        _para(doc, "", space=4)

    _para(doc, "Работа водителя и автомобиля", bold=True, space=2)
    _grid(
        doc,
        [
            [
                "Операция",
                "Время по графику",
                "Нулевой пробег, км",
                "Показание одометра, км",
                "Время фактическое",
            ],
            [
                "Выезд из гаража",
                _fmt_dt(wb.get("scheduled_departure_at")),
                _fmt_num(wb.get("zero_run_start_km"), 1),
                _fmt_num(wb.get("odometer_start_km"), 1),
                _fmt_dt(wb.get("departure_at")),
            ],
            [
                "Возвращение в гараж",
                _fmt_dt(wb.get("scheduled_return_at")),
                _fmt_num(wb.get("zero_run_end_km"), 1),
                _fmt_num(wb.get("odometer_end_km"), 1),
                _fmt_dt(wb.get("return_at")),
            ],
        ],
        [5.0, 4.6, 3.4, 4.6, 4.6],
        header=True,
    )
    _para(doc, "", space=4)

    _para(doc, "Движение горючего", bold=True, space=2)
    _pairs(doc, _fuel_rows(wb, cargo=True))
    _para(doc, "", space=4)

    _para(doc, "Задание водителю", bold=True, space=2)
    _grid(
        doc,
        [
            [
                "В чьё распоряжение (наименование и адрес заказчика)",
                "Время прибытия",
                "Адрес пункта погрузки",
                "Адрес пункта разгрузки",
                "Наименование груза",
                "Количество ездок",
                "Расстояние, км",
                "Перевезти, т",
            ],
            [
                " ".join(
                    x for x in (wb.get("customer_name", ""), wb.get("customer_address", "")) if x
                ),
                _fmt_dt(wb.get("planned_arrival_at")),
                wb.get("pickup_address", ""),
                wb.get("delivery_address", ""),
                wb.get("cargo_name", ""),
                str(wb.get("planned_trips") or ""),
                _fmt_num(wb.get("planned_distance_km"), 1),
                _fmt_num((wb.get("planned_cargo_kg") or 0) / 1000, 3)
                if wb.get("planned_cargo_kg")
                else "",
            ],
        ],
        [5.6, 2.6, 4.4, 4.4, 3.4, 1.8, 2.0, 1.8],
        header=True,
    )
    _para(doc, "", space=2)
    _para(doc, f"Выдать горючего {_fmt_num(wb.get('fuel_to_issue_l')) or _RULE} литров")
    _para(doc, "", space=4)

    _signature_line(
        doc,
        "Водительское удостоверение проверил, задание выдал. Диспетчер",
        wb.get("dispatcher_name", ""),
    )
    _medical_block(doc, wb)

    doc.add_page_break()
    _para(
        doc,
        "Оборотная сторона формы № 4-С",
        size=_SIZE_SMALL,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space=4,
    )
    _para(doc, "Последовательность выполнения задания", bold=True, space=2)
    rows = [
        [
            "Номер ездки",
            "Пункт погрузки, разгрузки и перецепки прицепов",
            "Номер прицепа прибывших",
            "Номер прицепа убывших",
            "Прибытие",
            "Убытие",
            "Номера ТТД",
            "Наименование грузоотправителя (грузополучателя)",
            "Перевезено, кг",
        ]
    ]
    for index, trip in enumerate(trips, start=1):
        rows.append(
            [
                str(trip.get("seq") or index),
                trip.get("place_to_name") or trip.get("destination_text") or "",
                trip.get("trailer_in", ""),
                trip.get("trailer_out", ""),
                _fmt_time(trip.get("returned_at")),
                _fmt_time(trip.get("departed_at")),
                trip.get("ttd_numbers", ""),
                trip.get("consignor", ""),
                str(trip.get("cargo_weight_kg") or ""),
            ]
        )
    for _ in range(max(0, 8 - len(trips))):
        rows.append([""] * 9)
    _grid(doc, rows, [1.4, 5.0, 2.0, 2.0, 1.8, 1.8, 3.0, 5.0, 2.0], header=True)

    _para(doc, "", space=4)
    _pairs(
        doc,
        [
            ("ТТД в количестве, шт", str(wb.get("ttd_count") or "")),
            ("Таксировка (прописью)", wb.get("taxation_text", "")),
        ],
    )
    _signature_line(doc, "ТТД сдал. Водитель", wb.get("ttd_handed_by", ""))
    _signature_line(doc, "ТТД принял. Диспетчер", wb.get("ttd_taken_by", ""))

    downtimes = wb.get("downtimes") or []
    if downtimes:
        _para(doc, "", space=4)
        _para(doc, "Простои на линии", bold=True, space=2)
        rows = [["Наименование", "Причина", "Начало", "Окончание", "Ответственное лицо"]]
        for item in downtimes:
            rows.append(
                [
                    item.get("name", ""),
                    item.get("reason", ""),
                    _fmt_dt(item.get("started_at")),
                    _fmt_dt(item.get("ended_at")),
                    item.get("responsible_name", ""),
                ]
            )
        _grid(doc, rows, [5.0, 5.0, 3.6, 3.6, 5.0], header=True)

    _para(doc, "", space=4)
    _para(doc, "Результаты работы автомобиля и прицепов", bold=True, space=2)
    _pairs(
        doc,
        [
            ("Время в наряде: всего", _fmt_minutes(wb.get("total_time_min"))),
            ("в том числе в движении", _fmt_minutes(wb.get("driving_time_min"))),
            ("в простое: автомобиля", _fmt_minutes(wb.get("idle_vehicle_min"))),
            ("под погрузкой, разгрузкой", _fmt_minutes(wb.get("idle_loading_min"))),
            ("по техническим неисправностям", _fmt_minutes(wb.get("idle_repair_min"))),
            ("сверхнормативный простой", _fmt_minutes(wb.get("idle_extra_min"))),
            ("Количество ездок", str(wb.get("trips_done") or "")),
            ("заездов", str(wb.get("stops_done") or "")),
            ("Пробег общий, км", _fmt_num(wb.get("run_total_km"), 1)),
            ("в том числе с грузом, км", _fmt_num(wb.get("run_loaded_km"), 1)),
            ("в том числе на прицепах, км", _fmt_num(wb.get("run_trailer_km"), 1)),
            ("Перевезено, кг", str(wb.get("cargo_done_kg") or "")),
            ("Выполнено, ткм", _fmt_num(wb.get("cargo_tkm"))),
            ("Автомобиль, дни в работе", str(wb.get("days_in_work") or "")),
            ("Зарплата автомобиля, руб. коп.", _fmt_num(wb.get("salary_vehicle_rub"))),
            ("Зарплата прицепа, руб. коп.", _fmt_num(wb.get("salary_trailer_rub"))),
            ("Итого, руб. коп.", _fmt_num(wb.get("salary_total_rub"))),
            ("Коды марок: автомобиля", wb.get("vehicle_code", "")),
        ],
    )
    _para(doc, "", space=4)
    _signature_line(doc, "Таксировщик", wb.get("taxer_name", ""))

    if wb.get("special_marks"):
        _para(doc, "Особые отметки", bold=True, space=2)
        _para(doc, wb["special_marks"])
    if wb.get("owner_marks"):
        _para(doc, "Отметки организации-владельца автотранспорта", bold=True, space=2)
        _para(doc, wb["owner_marks"])


def build_waybill_docx(waybill: dict, trips: list[dict], output_path: Path) -> Path:
    """Собирает печатную форму путевого листа — № 3 или 4-С по полю form."""
    doc = Document()
    _landscape(doc)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = _SIZE

    if waybill.get("form") == "4-S":
        _build_form4s(doc, waybill, trips)
    else:
        _build_form3(doc, waybill, trips)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
