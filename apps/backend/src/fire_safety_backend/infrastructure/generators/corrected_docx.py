"""Исправленный документ КОПИЕЙ оригинала, с сохранением форматирования.

Почему копия, а не новый файл из исправленного текста: пользователь этот
документ потом отправляет и подписывает. Собранный заново DOCX теряет всё —
заголовки, таблицы, шрифты, колонтитулы, бланк, — и на выходе получается
простыня текста, которую всё равно придётся переверстывать руками. Тогда
смысла в кнопке «скачать исправленный» нет.

Поэтому берём исходный файл и заменяем текст ВНУТРИ него, по парам
«было → стало» из результата проверки. Форматирование остаётся, потому что
мы не пересобираем документ, а правим содержимое существующих run'ов.

Для источников, которых нельзя «поправить внутри» (PDF, скан, вставленный
руками текст), собираем простой DOCX из исправленного текста — иначе
скачивать было бы нечего.
"""

from __future__ import annotations

import logging
import shutil
from typing import TYPE_CHECKING

from docx import Document

from ... import config
from .. import secure_files

if TYPE_CHECKING:
    from pathlib import Path

    from docx.text.paragraph import Paragraph

log = logging.getLogger(__name__)

SUFFIX = "_исправленный"


def _replacements(errors: list[dict]) -> list[tuple[str, str]]:
    """Пары «было → стало», пригодные для подстановки.

    Отбрасываем пустые и «замены на себя» — они ничего не меняют, но
    удлиняют проход. Сортируем по длине «было» по убыванию: иначе короткая
    замена может съесть часть более длинной (например «в» внутри «в течении»)
    и сломать её.
    """
    pairs: list[tuple[str, str]] = []
    for e in errors:
        if not isinstance(e, dict):
            continue
        before = str(e.get("before") or "").strip()
        after = str(e.get("after") or "").strip()
        if not before or before == after:
            continue
        pairs.append((before, after))
    # dict.fromkeys — убрать дубли, сохранив порядок первого появления.
    pairs = list(dict.fromkeys(pairs))
    pairs.sort(key=lambda p: len(p[0]), reverse=True)
    return pairs


def _apply_within_runs(paragraph: Paragraph, pairs: list[tuple[str, str]]) -> int:
    """Заменяет то, что целиком уложилось в один run, — оформление сохраняется.

    Это основной путь: жирный/курсив у исправленного слова остаётся жирным.
    Слияние run'ов (ниже) портит оформление, поэтому к нему прибегаем только
    для правок, которые действительно разрезаны между run'ами.
    """
    applied = 0
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        for before, after in pairs:
            if before in text:
                # Заменяем всё сразу: count не нужен, повторов в одном run мало.
                text = text.replace(before, after)
                applied += 1
        if text != run.text:
            run.text = text
    return applied


def _apply_across_runs(paragraph: Paragraph, pairs: list[tuple[str, str]]) -> int:
    """Правки, разрезанные границей run'ов.

    Word дробит абзац на run'ы по своей логике — из-за проверки орфографии,
    отслеживания правок, смены шрифта, — и слово запросто оказывается
    разорванным. Такую правку иначе не применить: приходится собрать текст
    абзаца целиком, заменить и положить обратно в первый run. Оформление
    остальных run'ов при этом теряется — поэтому путь и вторичный.
    """
    runs = paragraph.runs
    full = "".join(r.text for r in runs)
    applied = 0
    for before, after in pairs:
        if before in full:
            full = full.replace(before, after)
            applied += 1
    if not applied:
        return 0
    runs[0].text = full
    # Очищаем, но НЕ удаляем: удаление run'ов ломает нумерацию, гиперссылки
    # и закладки, привязанные к элементам абзаца.
    for r in runs[1:]:
        r.text = ""
    return applied


def _apply_to_paragraph(paragraph: Paragraph, pairs: list[tuple[str, str]]) -> int:
    """Заменяет текст в одном абзаце. Возвращает число замен."""
    runs = paragraph.runs
    if not runs or not "".join(r.text for r in runs).strip():
        return 0

    applied = _apply_within_runs(paragraph, pairs)

    # Осталось ли что-то, что видно только в склеенном тексте абзаца?
    full = "".join(r.text for r in runs)
    spanning = [(b, a) for b, a in pairs if b in full]
    if spanning:
        applied += _apply_across_runs(paragraph, spanning)
    return applied


def _iter_paragraphs(doc: Document) -> list[Paragraph]:
    """Абзацы тела + всех ячеек таблиц.

    В договорах и приказах существенная часть текста живёт в таблицах —
    без них правки прошли бы мимо половины документа.
    """
    out = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
    return out


def _plain_docx(corrected_text: str, dest: Path) -> Path:
    """Простой DOCX из исправленного текста — для источников без вёрстки."""
    doc = Document()
    for block in corrected_text.split("\n"):
        doc.add_paragraph(block)
    doc.save(str(dest))
    return dest


def _plain_result(corrected_text: str, dest: Path) -> tuple[Path, bool]:
    """Простой DOCX по логическому пути, с шифрованием на диске."""
    with secure_files.encrypted_output(dest) as writable:
        _plain_docx(corrected_text, writable)
    return dest, False


def build_corrected_docx(
    corrected_text: str,
    errors: list[dict],
    source_path: Path | None = None,
) -> tuple[Path, bool]:
    """Готовит исправленный документ и возвращает (путь, правки_в_копии).

    Второй элемент — True, если удалось поправить копию оригинала (значит
    форматирование на месте), и False, если документ собран заново из текста.
    Интерфейсу это нужно, чтобы честно сказать пользователю, что он получит.

    `source_path` — ЛОГИЧЕСКИЙ путь загруженного файла (без `.enc`): из него
    берутся имя результата и расширение, а сам файл читается через
    расшифрованную копию.
    """
    out_name = "документ"
    if source_path is not None:
        out_name = source_path.stem
    dest = config.OUTPUT_DIR / f"{out_name}{SUFFIX}.docx"

    is_docx = source_path is not None and source_path.suffix.lower() == ".docx"
    if not is_docx:
        return _plain_result(corrected_text, dest)

    pairs = _replacements(errors)
    try:
        with secure_files.encrypted_output(dest) as writable:
            # Копия делается с расшифрованной копии оригинала: на диске
            # оригинал лежит зашифрованным, а python-docx умеет только файл.
            with secure_files.plaintext(source_path) as readable:
                shutil.copyfile(readable, writable)
            doc = Document(str(writable))
            total = sum(_apply_to_paragraph(p, pairs) for p in _iter_paragraphs(doc))
            doc.save(str(writable))
    except secure_files.StorageUnprotected:
        # Не «нестандартный DOCX», а отказ хранилища — простой документ тут
        # тоже сохранить не удастся, и подменять причину в логе незачем.
        raise
    except Exception:
        # Порченный/нестандартный DOCX не должен лишать пользователя
        # исправленного текста — отдаём хотя бы простой документ.
        log.exception("Не удалось поправить копию %s — отдаю простой DOCX", source_path)
        return _plain_result(corrected_text, dest)

    log.info("Исправленный документ: %s (замен: %d)", dest.name, total)
    return dest, True
