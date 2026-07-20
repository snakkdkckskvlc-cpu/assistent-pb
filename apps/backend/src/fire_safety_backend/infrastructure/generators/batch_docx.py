"""Сводный DOCX-отчёт пакетной проверки договоров.

Обычный документ без фирменного бланка — это внутренний рабочий отчёт
для юриста/директора, а не исходящее письмо.
"""

from __future__ import annotations

from datetime import date
from typing import TYPE_CHECKING

from docx import Document

if TYPE_CHECKING:
    from pathlib import Path

_CRIT_ORDER = {"красный": 0, "жёлтый": 1, "желтый": 1, "зелёный": 2, "зеленый": 2}
_CRIT_MARK = {"красный": "🔴", "жёлтый": "🟡", "желтый": "🟡", "зелёный": "🟢", "зеленый": "🟢"}


def _finding_counts(findings: list[dict]) -> str:
    counts: dict[str, int] = {}
    for f in findings:
        crit = str(f.get("критичность", "")).casefold()
        counts[crit] = counts.get(crit, 0) + 1
    parts = []
    for crit in ("красный", "жёлтый", "зелёный"):
        n = counts.get(crit, 0) + (counts.get(crit.replace("ё", "е"), 0) if "ё" in crit else 0)
        if n:
            parts.append(f"{_CRIT_MARK[crit]} {crit}: {n}")
    return ", ".join(parts) if parts else "находок нет"


def build_batch_docx(items: list[dict], output_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Сводный отчёт по пакетной проверке договоров", level=0)
    doc.add_paragraph(f"Дата: {date.today().strftime('%d.%m.%Y')}. Файлов: {len(items)}.")

    for item in items:
        doc.add_heading(item.get("файл", "?"), level=1)

        if item.get("пропущен"):
            doc.add_paragraph(f"Тип: {item.get('тип', '?')}. {item.get('причина', 'Пропущен.')}")
            continue

        findings = item.get("находки") or []
        doc.add_paragraph(f"Тип: договор. Находок: {len(findings)} ({_finding_counts(findings)}).")

        for f in sorted(
            (f for f in findings if isinstance(f, dict)),
            key=lambda f: _CRIT_ORDER.get(str(f.get("критичность", "")).casefold(), 9),
        ):
            crit = str(f.get("критичность", "?"))
            head = doc.add_paragraph()
            head.add_run(f"{_CRIT_MARK.get(crit.casefold(), '•')} {crit.upper()}: ").bold = True
            head.add_run(str(f.get("в_чём_риск", "")))
            quote = str(f.get("цитата_из_договора", "")).strip()
            if quote:
                doc.add_paragraph(f"Цитата: «{quote}»")
            fix = str(f.get("предложение_правки", "")).strip()
            if fix:
                doc.add_paragraph(f"Предложение: {fix}")
            norm = str(f.get("ссылка_на_норму", "")).strip()
            if norm:
                doc.add_paragraph(f"Норма: {norm}")

        summary = item.get("сводка") or {}
        verdict = str(summary.get("общий_вывод", "")).strip()
        if verdict:
            p = doc.add_paragraph()
            p.add_run("Вывод: ").bold = True
            p.add_run(verdict)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
