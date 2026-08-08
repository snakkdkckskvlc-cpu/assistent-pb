"""Генерация DOCX-письма поверх фирменного бланка.

Ожидает, что в шаблоне letterhead.docx есть плейсхолдеры:
  {{date}}, {{recipient}}, {{subject}}, {{greeting}}, {{body}},
  {{sender_position}}, {{sender_name}}

Сам шаблон собирается из канонического бланка компании скриптом
scripts/build_letterhead_template.py — не редактируйте letterhead.docx
руками, правьте исходный «бланк ПожСервис …docx» и пересоберите.

Многострочные значения ({{recipient}} = «Директору …\\nИванову А.А.»,
многоабзацное {{body}}) корректно разбиваются на параграфы с сохранением
формата исходного плейсхолдер-параграфа.

Если шаблона нет — генерирует чистый DOCX без бланка.
"""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from datetime import date
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ... import config

log = logging.getLogger(__name__)

if TYPE_CHECKING:
    from pathlib import Path

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def letterhead_requisites() -> list[str] | None:
    """Строки реквизитов из самого бланка. None — шаблона нет.

    Предпросмотр письма показывал реквизиты СВОЕЙ копией, зашитой строкой в
    letter.html. Две копии одного и того же расходятся при первой же правке
    реквизитов, а расходятся они в худшем месте: на экране человек видит один
    расчётный счёт, в отправленном документе — другой. Подмена счёта в
    исходящем письме — классическая схема мошенничества, и выглядеть она будет
    как наша ошибка.

    Поэтому источник один — тот файл, который и уходит контрагенту. Реквизиты
    лежат в первой ячейке шапки-таблицы, рядом с плейсхолдером {{recipient}};
    плейсхолдеры и пустые строки отбрасываются.
    """
    if not config.LETTERHEAD_TEMPLATE.exists():
        return None
    try:
        doc = Document(str(config.LETTERHEAD_TEMPLATE))
        if not doc.tables:
            return None
        cell = doc.tables[0].rows[0].cells[0]
        lines = [p.text.strip() for p in cell.paragraphs]
        return [ln for ln in lines if ln and not _PLACEHOLDER_RE.search(ln)]
    except Exception:
        # Бланк может быть заменён организацией на свой (см. security.md), и
        # структура окажется другой. Это не повод ронять страницу письма:
        # предпросмотр просто скажет, что реквизиты не прочитались.
        log.exception("Не удалось прочитать реквизиты из бланка")
        return None


def _substitute(text: str, mapping: dict[str, str]) -> str | None:
    """Заменяет все плейсхолдеры в исходном тексте параграфа за один проход.

    Один re.sub по оригинальному тексту (а не последовательные .replace()
    по ключам) исключает повторную подстановку: значение одного ключа не
    может случайно содержать «{{другой_ключ}}» и быть заменено ещё раз —
    re.sub никогда не пересканирует уже подставленный текст.
    Возвращает None, если в тексте не было ни одного плейсхолдера из mapping
    (сигнал вызывающему коду ничего не менять).
    """
    found = False

    def _repl(m: re.Match[str]) -> str:
        nonlocal found
        key = m.group(1)
        if key not in mapping:
            return m.group(0)
        found = True
        return str(mapping[key]) if mapping[key] else ""

    new_text = _PLACEHOLDER_RE.sub(_repl, text)
    return new_text if found else None


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    new_text = _substitute(paragraph.text, mapping)
    if new_text is None:
        return

    lines = new_text.split("\n") if new_text else [""]
    first_line = lines[0]

    # Пересобираем runs так, чтобы сохранить формат первого run
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = first_line
    else:
        paragraph.add_run(first_line)

    # Остальные строки — как отдельные параграфы сразу после текущего,
    # с копией формата (стили, выравнивание, шрифт).
    prev_p = paragraph
    for line in lines[1:]:
        new_p_xml = deepcopy(paragraph._p)
        prev_p._p.addnext(new_p_xml)
        # Нужен объект Paragraph, а не только XML — получим через parent
        from docx.text.paragraph import Paragraph as _P

        new_p = _P(new_p_xml, paragraph._parent)
        for r in new_p.runs:
            r.text = ""
        if new_p.runs:
            new_p.runs[0].text = line
        else:
            new_p.add_run(line)
        prev_p = new_p


def _apply_mapping(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def build_letter_docx(letter: dict, output_path: Path) -> Path:
    """Собирает DOCX-письмо. Если фирменный бланк есть — использует его."""
    mapping = {
        "date": date.today().strftime("%d.%m.%Y"),
        "recipient": letter.get("получатель", ""),
        "subject": letter.get("тема", ""),
        "greeting": letter.get("обращение", ""),
        "body": letter.get("тело", ""),
        "sender_position": letter.get("должность_отправителя_placeholder", "[должность]"),
        "sender_name": letter.get("фио_отправителя_placeholder", "[Фамилия И.О.]"),
    }

    if config.LETTERHEAD_TEMPLATE.exists():
        doc = Document(str(config.LETTERHEAD_TEMPLATE))
        _apply_mapping(doc, mapping)
    else:
        # Отсутствие бланка — НЕ мелочь: письмо уходит контрагенту без
        # реквизитов, ИНН и банковских данных, то есть как обычный текст, а
        # не как официальный документ компании. Отказ при этом внешне
        # незаметен — файл создаётся, скачивается и открывается. Такое уже
        # случилось: шаблон пропал, и разницу заметили только сравнив два
        # письма вручную. Поэтому пишем в лог ошибкой и помечаем результат,
        # чтобы интерфейс сказал об этом пользователю ДО отправки.
        log.error(
            "Шаблон бланка не найден (%s) — письмо будет БЕЗ реквизитов компании. "
            "Восстановить: python scripts/build_letterhead_template.py --source <бланк.docx>",
            config.LETTERHEAD_TEMPLATE,
        )
        # Fallback — минимальный документ
        doc = Document()
        doc.add_paragraph(mapping["date"]).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph()
        if mapping["recipient"]:
            doc.add_paragraph(mapping["recipient"])
        if mapping["subject"]:
            p = doc.add_paragraph()
            p.add_run(f"Касательно: {mapping['subject']}").bold = True
        doc.add_paragraph()
        doc.add_paragraph(mapping["greeting"])
        doc.add_paragraph()
        for para in mapping["body"].split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()
        doc.add_paragraph("С уважением,")
        doc.add_paragraph(mapping["sender_position"])
        doc.add_paragraph(mapping["sender_name"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
