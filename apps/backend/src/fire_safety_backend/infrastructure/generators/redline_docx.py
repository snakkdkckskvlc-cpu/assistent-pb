"""Договор с правками в режиме рецензирования Word.

Зачем отдельно от legal_docx. Тот собирает ОТЧЁТ о разборе — его читают перед
переговорами. Этот собирает сам ДОГОВОР, в котором предложенные формулировки
внесены как правки: контрагент открывает файл в Word и принимает или
отклоняет каждую кнопкой. До этого юрист переносил формулировки из списка
находок руками, по одной.

Почему настоящие `w:ins`/`w:del`, а не зачёркнутый и цветной текст. Визуальная
разметка выглядит как редлайн, но ею нельзя воспользоваться: у второй стороны
нет кнопки «принять», сравнение версий в Word её не видит, и при копировании
в свой шаблон зачёркнутое едет вместе с текстом. Правка в режиме рецензирования
— это стандартный способ обмена договорами, и юрист ждёт именно его.

Правки вносятся ТОЛЬКО там, где цитата дословно найдена в оригинале
(`_цитата_найдена`) и модель предложила замену. Всё остальное молча
пропускается: подставить правку не в то место в договоре хуже, чем не
подставить вовсе.
"""

from __future__ import annotations

import copy
import logging
import shutil
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ... import config
from ...services.uploads import original_name
from .. import secure_files

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

log = logging.getLogger(__name__)

SUFFIX = "_с_правками"

# Имя в плашке правки у контрагента. Не «Claude» и не название модели:
# получатель договора не должен гадать, кто такой Claude, а компания отвечает
# за содержание правок сама.
AUTHOR = "Ассистент ПБ"

# Дата правок. Word показывает её во всплывающей подсказке. Формат — ISO 8601
# с Z, другого Word не понимает.
_DATE_FMT = "%Y-%m-%dT%H:%M:%SZ"

# Цитата короче этого в договоре почти наверняка встречается не один раз, и
# правка уедет не в тот пункт. Замерено на текстах договоров: «Подрядчик»
# встречается десятки раз, «оплата» — единицы. 40 символов — та длина, начиная
# с которой совпадение перестаёт быть случайным.
MIN_QUOTE_LEN = 40


def _iter_paragraphs(doc: Document) -> list[Paragraph]:
    """Абзацы тела и всех ячеек таблиц.

    В договорах существенная часть условий живёт в таблицах — без них правки
    прошли бы мимо приложений со сроками и ценами. Та же причина, что в
    corrected_docx._iter_paragraphs.
    """
    out = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
    return out


def _run_text_element(run_el) -> object | None:
    """Текстовый узел run'а. None — значит run несёт разрыв, табуляцию или
    картинку, и резать его по символам нельзя."""
    node = run_el.find(qn("w:t"))
    return node


def _clone_run_with_text(run_el, text: str):
    """Копия run'а с тем же оформлением и другим текстом."""
    clone = copy.deepcopy(run_el)
    node = clone.find(qn("w:t"))
    if node is None:
        node = OxmlElement("w:t")
        clone.append(node)
    node.text = text
    # Без xml:space="preserve" Word съедает пробелы по краям, и слова
    # склеиваются: «пенюа0,1%» вместо «пеню 0,1 %».
    node.set(qn("xml:space"), "preserve")
    return clone


def _as_deleted_run(run_el, text: str):
    """Run для <w:del>: текст переезжает из w:t в w:delText.

    Word не покажет удаление, если оставить w:t — элемент внутри w:del обязан
    быть именно delText, иначе файл считается повреждённым и Word молча
    выбрасывает правку при открытии.
    """
    clone = copy.deepcopy(run_el)
    for node in clone.findall(qn("w:t")):
        clone.remove(node)
    node = OxmlElement("w:delText")
    node.text = text
    node.set(qn("xml:space"), "preserve")
    clone.append(node)
    return clone


def _wrap(tag: str, change_id: int, when: str, *runs):
    """<w:ins> или <w:del> вокруг готовых run'ов."""
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(change_id))
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), when)
    for r in runs:
        el.append(r)
    return el


def _mark_quote(
    paragraph: Paragraph, quote: str, replacement: str, when: str, first_id: int
) -> int:
    """Помечает цитату в абзаце как удалённую, а замену — как вставленную.

    Возвращает число израсходованных идентификаторов правки (0 — цитата в этом
    абзаце не найдена). Каждая правка Word — это отдельный w:id, и повтор
    идентификатора склеивает разные правки в одну.
    """
    runs = paragraph.runs
    if not runs:
        return 0
    full = "".join(r.text for r in runs)
    start = full.find(quote)
    if start < 0:
        return 0
    end = start + len(quote)

    # Раскладка абзаца по run'ам: для каждого — его границы в общем тексте.
    spans: list[tuple[object, int, int]] = []
    pos = 0
    for run in runs:
        spans.append((run._element, pos, pos + len(run.text)))
        pos += len(run.text)

    touched = [(el, s, e) for el, s, e in spans if s < end and e > start and e > s]
    if not touched:
        return 0
    # Run с картинкой или разрывом внутри цитаты резать нельзя — пропускаем
    # правку целиком, чтобы не потерять содержимое.
    if any(_run_text_element(el) is None for el, _, _ in touched):
        log.info("redline: цитата пересекает нетекстовый элемент, правка пропущена")
        return 0

    first_el, first_start, _ = touched[0]
    last_el, _, last_end = touched[-1]
    parent = first_el.getparent()
    anchor = first_el

    # Хвост оригинала ПОСЛЕ цитаты — вставляем сразу, пока не удалили run'ы.
    tail = full[end:last_end] if last_end > end else ""
    # Голова ДО цитаты остаётся обычным текстом в своём оформлении.
    head = full[first_start:start] if start > first_start else ""

    new_nodes = []
    if head:
        new_nodes.append(_clone_run_with_text(first_el, head))

    deleted = _as_deleted_run(first_el, quote)
    new_nodes.append(_wrap("w:del", first_id, when, deleted))

    inserted = _clone_run_with_text(first_el, replacement)
    new_nodes.append(_wrap("w:ins", first_id + 1, when, inserted))

    if tail:
        new_nodes.append(_clone_run_with_text(last_el, tail))

    for node in new_nodes:
        parent.insert(parent.index(anchor), node)
    for el, _, _ in touched:
        parent.remove(el)
    return 2


def _usable_findings(findings: list[dict]) -> list[tuple[str, str]]:
    """Пары «цитата → предложенная формулировка», пригодные для правки.

    Отсекается всё, где подстановка была бы опасной:
      * цитата не найдена в оригинале дословно (`_цитата_найдена` False) —
        значит модель её перефразировала, и куда её ставить, неизвестно;
      * пустое предложение правки — находка есть, формулировки нет;
      * замена совпадает с цитатой — правка ничего не меняет, но создаёт в
        Word пустую пометку, которую контрагент обязан просмотреть;
      * слишком короткая цитата — см. MIN_QUOTE_LEN.

    Длинные цитаты идут первыми: если одна цитата содержится в другой,
    правка по длинной должна примениться раньше, иначе короткая разрежет
    длинную пополам и вторая правка не найдёт своего места.
    """
    pairs: list[tuple[str, str]] = []
    for f in findings:
        if not isinstance(f, dict):
            continue
        if f.get("_цитата_найдена") is False:
            continue
        quote = str(f.get("цитата_из_договора") or "").strip()
        fix = str(f.get("предложение_правки") or "").strip()
        if not quote or not fix or quote == fix:
            continue
        if len(quote) < MIN_QUOTE_LEN:
            log.info("redline: цитата короче %d символов, пропущена", MIN_QUOTE_LEN)
            continue
        pairs.append((quote, fix))
    pairs = list(dict.fromkeys(pairs))
    pairs.sort(key=lambda p: len(p[0]), reverse=True)
    return pairs


def build_redline_docx(
    findings: list[dict],
    source_path: Path | None,
    when: str,
) -> tuple[Path, int, int]:
    """Готовит договор с правками. Возвращает (путь, внесено, всего_пригодных).

    Разница между «внесено» и «всего» — правки, чью цитату не удалось найти в
    самом DOCX, хотя в извлечённом тексте она была. Так бывает, когда цитата
    пересекает границу абзаца или таблицы. Интерфейс обязан показать эту
    разницу: юрист должен знать, что часть правок придётся внести руками.

    `when` передаётся снаружи, а не берётся из datetime.now(), чтобы результат
    был воспроизводим в тестах.
    """
    out_name = "договор"
    if source_path is not None:
        out_name = Path(original_name(source_path)).stem
    dest = config.OUTPUT_DIR / f"{out_name}{SUFFIX}.docx"

    pairs = _usable_findings(findings)
    if source_path is None or source_path.suffix.lower() != ".docx":
        # Из PDF и сканов исходной вёрстки нет, а редлайн без оригинала
        # бессмыслен: править нечего. Честнее не отдавать файл совсем, чем
        # отдать пересобранную простыню, выдав её за договор.
        raise ValueError("Правки в режиме рецензирования доступны только для файлов DOCX")

    applied = 0
    with secure_files.encrypted_output(dest) as writable:
        with secure_files.plaintext(source_path) as readable:
            shutil.copyfile(readable, writable)
        doc = Document(str(writable))
        paragraphs = _iter_paragraphs(doc)
        change_id = 1
        for quote, fix in pairs:
            for paragraph in paragraphs:
                used = _mark_quote(paragraph, quote, fix, when, change_id)
                if used:
                    change_id += used
                    applied += 1
                    break
        doc.save(str(writable))
    return dest, applied, len(pairs)
