"""DOCX с разбором одного договора.

Внутренний рабочий документ для юриста и директора, а не исходящее письмо, —
поэтому без фирменного бланка (та же логика, что в batch_docx).

Зачем он вообще. У проверки орфографии выгрузка есть, у пакетной проверки —
сводный отчёт, а разбор одиночного договора выгружался только копированием
руками. Между тем именно его несут на переговоры с контрагентом.

Порядок здесь тот же, что на экране: вывод первым, находки следом. Читающему
нужен ответ «подписывать или нет», а не двадцать пунктов до него.
"""

from __future__ import annotations

from datetime import date
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt

if TYPE_CHECKING:
    from pathlib import Path

_CRIT_ORDER = {"красный": 0, "жёлтый": 1, "желтый": 1, "зелёный": 2, "зеленый": 2}
# Словами, а не цветными кружками: кружок в Word рисуется шрифтом системы и на
# чужой машине выглядит иначе, а документ уходит контрагенту.
_CRIT_WORD = {
    "красный": "КРИТИЧНО",
    "жёлтый": "СУЩЕСТВЕННО",
    "желтый": "СУЩЕСТВЕННО",
    "зелёный": "ВНИМАНИЕ",
    "зеленый": "ВНИМАНИЕ",
}


def _crit_key(finding: dict) -> str:
    return str(finding.get("критичность", "")).casefold()


def _counts(findings: list[dict]) -> str:
    counts: dict[str, int] = {}
    for f in findings:
        counts[_crit_key(f)] = counts.get(_crit_key(f), 0) + 1
    parts = []
    for crit in ("красный", "жёлтый", "зелёный"):
        # Модель пишет и «жёлтый», и «желтый» — считаем оба написания. Условие
        # обязательно: у «красного» буквы «ё» нет, и без него замена возвращала
        # то же слово, а находка попадала в сумму дважды.
        alt = crit.replace("ё", "е")
        n = counts.get(crit, 0) + (counts.get(alt, 0) if alt != crit else 0)
        if n:
            parts.append(f"{_CRIT_WORD[crit].lower()}: {n}")
    return ", ".join(parts) if parts else "находок нет"


def build_legal_docx(result: dict, output_path: Path, *, source_name: str = "") -> Path:
    findings = [f for f in (result.get("находки") or []) if isinstance(f, dict)]
    summary = result.get("сводка") or {}

    doc = Document()
    doc.add_heading("Разбор договора", level=0)

    head = f"Дата разбора: {date.today().strftime('%d.%m.%Y')}."
    if source_name:
        head = f"Документ: {source_name}. " + head
    doc.add_paragraph(f"{head} Находок: {len(findings)} ({_counts(findings)}).")

    # Вывод — до находок. На экране он тоже наверху; документ, в котором
    # порядок другой, читался бы как другой документ.
    verdict = str(summary.get("общий_вывод", "")).strip()
    if verdict:
        p = doc.add_paragraph()
        p.add_run("Вывод: ").bold = True
        p.add_run(verdict)

    pros = [str(x) for x in (summary.get("плюсы_для_компании") or [])]
    cons = [str(x) for x in (summary.get("минусы_для_компании") or [])]
    if pros or cons:
        doc.add_heading("Сводка", level=1)
        for title, items in (("Плюсы для компании", pros), ("Минусы для компании", cons)):
            if not items:
                continue
            doc.add_paragraph(title).runs[0].bold = True
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

    if findings:
        doc.add_heading("Находки", level=1)
    for f in sorted(findings, key=lambda f: _CRIT_ORDER.get(_crit_key(f), 9)):
        crit = _CRIT_WORD.get(_crit_key(f), str(f.get("критичность", "")).upper() or "НАХОДКА")
        head_p = doc.add_paragraph()
        head_p.add_run(f"{crit}: ").bold = True
        head_p.add_run(str(f.get("в_чём_риск", "")))

        # Причина автоматического повышения уровня — это утверждение программы,
        # а не модели, и юрист должен видеть, с чем именно ему предлагают
        # согласиться.
        raised = str(f.get("_критичность_повышена", "")).strip()
        if raised:
            note = doc.add_paragraph(f"Уровень поднят расчётом: {raised}")
            note.runs[0].font.size = Pt(9)

        quote = str(f.get("цитата_из_договора", "")).strip()
        if quote:
            doc.add_paragraph(f"Цитата: «{quote}»")
            # Молчаливая деградация — худший отказ в этом проекте: если цитата
            # не нашлась в документе дословно, читающий обязан знать об этом до
            # того, как понесёт её контрагенту.
            if f.get("_цитата_найдена") is False:
                warn = doc.add_paragraph(
                    "Внимание: точная цитата в тексте документа не найдена — "
                    "возможно, модель её перефразировала. Сверьте с оригиналом."
                )
                warn.runs[0].font.size = Pt(9)

        fix = str(f.get("предложение_правки", "")).strip()
        if fix:
            doc.add_paragraph(f"Предложение правки: {fix}")

        norm = str(f.get("ссылка_на_норму", "")).strip()
        if norm:
            status = str(f.get("_норма_статус", ""))
            if status == "подтверждена":
                tail = " (статья найдена в переданных нормах)"
            elif status == "не_в_контексте":
                tail = " (этой статьи не было в переданных нормах — проверьте номер)"
            else:
                tail = " (из общих знаний модели, контекстом не подтверждена)"
            small = doc.add_paragraph(f"Норма: {norm}{tail}")
            small.runs[0].font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
